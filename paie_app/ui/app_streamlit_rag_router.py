# =======================================================================
# PAIE – Offline Personalized AI + Analytics
# Streamlit UI (single-file)
# =======================================================================

# ===== SECTION: Imports & Globals ======================================
import os
import sys
import json
import subprocess
import requests
import sqlite3
from pathlib import Path

import yaml
import streamlit as st
# must be the first Streamlit command
if "PAGE_CONFIGURED" not in st.session_state:
    st.set_page_config(page_title="PAIE", page_icon="🛡️", layout="wide")
    st.session_state["PAGE_CONFIGURED"] = True
import pandas as pd
import plotly.express as px

# Optional libs (DOCX/PDF exports)
HAS_DOCX = True
try:
    from docx import Document
except Exception:
    HAS_DOCX = False

HAS_REPORTLAB = True
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
except Exception:
    HAS_REPORTLAB = False

# Hide the Export BI tab (4th tab) by CSS toggle (PAIE_HIDE_BI=1/0)
if os.getenv("PAIE_HIDE_BI", "1") == "1":
    st.markdown(
        """
        <style>
          /* Hide the 4th tab button (Export BI) */
          div[data-baseweb="tab-list"] button:nth-of-type(4){display:none!important;}
        </style>
        """,
        unsafe_allow_html=True,
    )

# Wire core
sys.path.append(str(Path(__file__).resolve().parents[1]))
from core.router import Router  # noqa: E402

APP_DIR        = Path(__file__).resolve().parents[1]
SETTINGS       = APP_DIR / "config" / "settings.yaml"
DB_PATH        = APP_DIR / "paie.db"
PROFILES_DIR   = APP_DIR / "config" / "profiles"
TEMPLATES_DIR  = APP_DIR / "config" / "templates"
USERNAME       = "default"
HEADLINE       = "PAIE – Personalized AI Engine"  # <- change here if you want

# ===== SECTION: Page Config & Session Keys =============================
st.title(HEADLINE)

# Session init (needed for Clear input)
if "prompt_input" not in st.session_state:
    st.session_state["prompt_input"] = ""
if "__last_prompt" not in st.session_state:
    st.session_state["__last_prompt"] = ""

# ===== SECTION: Helper Functions ======================================
def list_ollama_models(host: str) -> list[str]:
    """Discover available models from Ollama REST, fall back to `ollama list`."""
    names = []
    try:
        r = requests.get(host.rstrip("/") + "/api/tags", timeout=3)
        r.raise_for_status()
        data = r.json() or {}
        for m in data.get("models", []):
            name = (m.get("name") or m.get("model") or "").strip()
            if name:
                names.append(name)
    except Exception:
        try:
            out = subprocess.check_output(["ollama", "list"], text=True, timeout=4)
            lines = [l.strip() for l in out.splitlines() if l.strip()]
            names = [ln.split()[0] for ln in lines if not ln.upper().startswith("NAME")]
        except Exception:
            pass
    seen, out = set(), []
    for n in names:
        if n not in seen:
            out.append(n)
            seen.add(n)
    return out

def load_profiles() -> dict:
    items = {}
    if PROFILES_DIR.exists():
        for p in PROFILES_DIR.glob("*.y*ml"):
            try:
                items[p.stem] = yaml.safe_load(p.read_text(encoding="utf-8")) or {}
            except Exception:
                pass
    if not items:
        items["default"] = {
            "name": "default",
            "tone": "neutral",
            "verbosity": "medium",
            "persona": "Helpful assistant",
            "style_guide": [],
            "domain_glossary": [],
        }
    return items

def load_templates() -> dict[str, str]:
    items = {"(none)": ""}
    if TEMPLATES_DIR.exists():
        for p in TEMPLATES_DIR.glob("*.txt"):
            try:
                items[p.stem] = p.read_text(encoding="utf-8")
            except Exception:
                pass
    return items

def ensure_columns():
    """Ensure new columns exist on interactions table."""
    con = sqlite3.connect(str(DB_PATH))
    cur = con.cursor()
    for coldef in [
        ("model_name", "TEXT"),
        ("profile_name", "TEXT"),
        ("template_name", "TEXT"),
        ("rating", "INTEGER"),
    ]:
        col, typ = coldef
        try:
            cur.execute(f"ALTER TABLE interactions ADD COLUMN {col} {typ}")
            con.commit()
        except Exception:
            pass  # already exists
    con.close()

def update_last_interaction_meta(username, model, profile_name, template_name):
    con = sqlite3.connect(str(DB_PATH))
    cur = con.cursor()
    row = cur.execute(
        """
        SELECT i.id FROM interactions i
        JOIN users u ON u.id = i.user_id
        WHERE u.username=? ORDER BY i.id DESC LIMIT 1
        """,
        (username,),
    ).fetchone()
    if row:
        cur.execute(
            "UPDATE interactions SET model_name=?, profile_name=?, template_name=? WHERE id=?",
            (model, profile_name, template_name, row[0]),
        )
        con.commit()
    con.close()

def set_last_rating(username, val):
    con = sqlite3.connect(str(DB_PATH))
    cur = con.cursor()
    row = cur.execute(
        """
        SELECT i.id FROM interactions i
        JOIN users u ON u.id = i.user_id
        WHERE u.username=? ORDER BY i.id DESC LIMIT 1
        """,
        (username,),
    ).fetchone()
    if row:
        cur.execute("UPDATE interactions SET rating=? WHERE id=?", (val, row[0]))
        con.commit()
    con.close()
    return row[0] if row else None

def build_profile_text(p: dict) -> str:
    """Turn YAML profile into a compact system hint."""
    lines = []
    if p.get("persona"):
        lines.append(f"Persona: {p['persona']}")
    if p.get("tone"):
        lines.append(f"Tone: {p['tone']}")
    if p.get("verbosity"):
        lines.append(f"Verbosity: {p['verbosity']}")
    sg = p.get("style_guide") or []
    if sg:
        lines.append("Style guide: " + "; ".join(sg))
    gl = p.get("domain_glossary") or []
    if gl:
        lines.append("Domain glossary: " + "; ".join(gl))
    return "[PROFILE]\n" + "\n".join(lines) + "\n[/PROFILE]\n"

def make_docx(prompt, response, model, structure):
    if not HAS_DOCX:
        return None
    from io import BytesIO
    doc = Document()
    doc.add_heading("PAIE Output", level=1)
    doc.add_paragraph(f"Model: {model}")
    doc.add_paragraph(f"Structure: {structure or '(none)'}")
    doc.add_heading("Prompt", level=2);   doc.add_paragraph(prompt)
    doc.add_heading("Response", level=2); doc.add_paragraph(response)
    buf = BytesIO(); doc.save(buf); buf.seek(0)
    return buf.getvalue()

def make_pdf(text):
    if not HAS_REPORTLAB:
        return None
    from io import BytesIO
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4
    x, y = 40, height - 40
    for line in text.splitlines():
        if y < 40:
            c.showPage()
            y = height - 40
        c.drawString(x, y, line[:120])
        y -= 14
    c.save()
    buf.seek(0)
    return buf.getvalue()

# ===== SECTION: Sidebar – Session ======================================
st.sidebar.header("Session")

with open(SETTINGS, "r", encoding="utf-8") as f:
    default_settings = yaml.safe_load(f)

available_models = list_ollama_models(default_settings["model"]["host"])
default_model    = default_settings["model"]["name"]
if not available_models:
    available_models = [default_model]
if default_model in available_models:
    available_models = [default_model] + [m for m in available_models if m != default_model]

model   = st.sidebar.selectbox("Model", options=available_models, index=0)
kind    = st.sidebar.selectbox("Structure", ["(none)", "user_story", "use_case", "test_case", "summary"], index=0)

profiles      = load_profiles()
profile_names = list(profiles.keys())
profile_name  = st.sidebar.selectbox(
    "Profile",
    options=profile_names,
    index=profile_names.index("default") if "default" in profile_names else 0,
)

templates     = load_templates()
template_name = st.sidebar.selectbox("Template", options=list(templates.keys()), index=0)

# Init DB + core router
ensure_columns()
r = Router(str(SETTINGS))
r.settings["model"]["name"] = model  # propagate chosen model

# ===== SECTION: Tabs ====================================================
tabs = st.tabs(["Generate", "History", "Analytics"])

# -----------------------------------------------------------------------
# SECTION: Generate
# -----------------------------------------------------------------------
with tabs[0]:
    # Input area (keep the same key; we only clear it via callback)
    prompt = st.text_area(
        "Prompt",
        key="prompt_input",
        height=160,
        placeholder="Ask for a User Story / Use Case / Test Case ...",
    )

    # Buttons row: Generate on left, Clear on far right
    left, spacer, right = st.columns([1, 6, 1])

    def _clear_prompt():
        st.session_state["prompt_input"] = ""
        st.session_state["__last_prompt"] = ""

    with left:
        run_clicked = st.button("Generate", type="primary", key="btn_generate")

    with right:
        st.button("Clear input", key="btn_clear", on_click=_clear_prompt)

    # Main action
    if run_clicked and st.session_state.get("prompt_input", "").strip():
        prompt = st.session_state["prompt_input"]
        prof_txt = build_profile_text(profiles[profile_name])
        tpl_txt  = templates.get(template_name, "")
        final_prompt = (
            f"{prof_txt}"
            f"{('[TEMPLATE]\\n' + tpl_txt + '\\n[/TEMPLATE]\\n') if tpl_txt else ''}"
            f"{prompt}"
        )

        with st.spinner("Thinking locally..."):
            k = None if kind == "(none)" else kind
            out, ms = r.run(username=USERNAME, prompt=final_prompt, kind=k)

        # Save meta
        update_last_interaction_meta(USERNAME, model, profile_name, template_name)

        st.success(f"Done in {ms} ms")
        st.markdown(out)

        # Downloads
        md = f"# Model: {model}\n# Structure: {k}\n\n# Prompt\n\n{prompt}\n\n# Response\n\n{out}"
        st.download_button("Download Markdown", data=md, file_name="output.md", mime="text/markdown")

        if HAS_DOCX:
            docx_bytes = make_docx(prompt, out, model, k)
            if docx_bytes:
                st.download_button(
                    "Download DOCX",
                    data=docx_bytes,
                    file_name="output.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
        else:
            st.info("DOCX export requires python-docx (optional).")

        if HAS_REPORTLAB:
            pdf_bytes = make_pdf(f"Model: {model}\nStructure: {k}\n\nPrompt:\n{prompt}\n\nResponse:\n{out}")
            if pdf_bytes:
                st.download_button("Download PDF", data=pdf_bytes, file_name="output.pdf", mime="application/pdf")
        else:
            st.info("PDF export requires reportlab (optional).")

        # Feedback
        st.write("Feedback:")
        b1, b2 = st.columns(2)
        if b1.button("👍 Good"):
            _id = set_last_rating(USERNAME, 1)
            st.toast(f"Thanks! Saved 👍 for interaction #{_id}")
        if b2.button("👎 Needs work"):
            _id = set_last_rating(USERNAME, -1)
            st.toast(f"Saved 👎 for interaction #{_id}")

# -----------------------------------------------------------------------
# SECTION: History
# -----------------------------------------------------------------------
with tabs[1]:
    st.subheader("History")
    con = sqlite3.connect(str(DB_PATH))
    q = st.text_input("Search", placeholder="filter by keyword (optional)")

    sql = """
      SELECT i.id, i.created_at, i.structure_kind, i.prompt, i.response,
             i.model_name, i.profile_name, i.template_name, i.rating
      FROM interactions i
      JOIN users u ON u.id = i.user_id
      WHERE u.username = ?
    """
    params = [USERNAME]
    if q:
        sql += " AND (i.prompt LIKE ? OR i.response LIKE ?)"
        params += [f"%{q}%", f"%{q}%"]
    sql += " ORDER BY i.id DESC LIMIT 200"
    df = pd.read_sql_query(sql, con, params=params)

    with st.expander("🧹 Clear history"):
        st.caption("Permanently deletes saved prompts/responses.")
        if st.button("Delete all history"):
            uid = con.execute("SELECT id FROM users WHERE username=?", (USERNAME,)).fetchone()
            if uid:
                con.execute("DELETE FROM interactions WHERE user_id=?", (uid[0],))
                con.commit()
                st.success("History cleared. Click ↻ Rerun (top-right) to refresh.")
            else:
                st.info("No records found.")
    con.close()

    if df.empty:
        st.info("No history yet.")
    else:
        df["date"] = pd.to_datetime(df["created_at"]).dt.date
        for date, grp in df.groupby("date", sort=False):
            st.markdown(f"### {date}")
            for _, row in grp.iterrows():
                header = (
                    f"#{row['id']}  •  {row['created_at']}  •  "
                    f"{row['structure_kind'] or '(none)'}  •  {row['model_name'] or model}"
                )
                with st.expander(header):
                    st.markdown(
                        f"**Profile**: {row['profile_name'] or profile_name}  |  "
                        f"**Template**: {row['template_name'] or '(none)'}  |  "
                        f"**Rating**: {row['rating'] if pd.notna(row['rating']) else '—'}"
                    )
                    st.markdown("**Prompt**")
                    st.code(row["prompt"])
                    st.markdown("**Response**")
                    st.markdown(row["response"])

                    md_hist = f"# Prompt\n\n{row['prompt']}\n\n# Response\n\n{row['response']}"
                    st.download_button(
                        "Download as Markdown",
                        data=md_hist,
                        file_name=f"paie_{row['id']}.md",
                        key=f"md-{row['id']}",
                    )
                    if HAS_DOCX:
                        docx_bytes = make_docx(
                            row["prompt"],
                            row["response"],
                            row["model_name"] or model,
                            row["structure_kind"],
                        )
                        if docx_bytes:
                            st.download_button(
                                "Download as DOCX",
                                data=docx_bytes,
                                file_name=f"paie_{row['id']}.docx",
                                key=f"docx-{row['id']}",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            )

# -----------------------------------------------------------------------
# SECTION: Analytics
# -----------------------------------------------------------------------
with tabs[2]:
    st.subheader("Usage & Performance (from SQLite)")
    con = sqlite3.connect(str(DB_PATH))
    df = pd.read_sql_query(
        "SELECT created_at, structure_kind, latency_ms, model_name, rating FROM interactions ORDER BY id DESC LIMIT 2000",
        con,
    )
    con.close()

    if len(df) == 0:
        st.info("No data yet. Generate something first.")
    else:
        df["date"] = pd.to_datetime(df["created_at"]).dt.date
        c1, c2 = st.columns(2)
        c1.plotly_chart(px.bar(df, x="structure_kind", title="Structure usage count"), use_container_width=True)
        c2.plotly_chart(
            px.line(df.groupby("date")["latency_ms"].mean().reset_index(), x="date", y="latency_ms", title="Avg latency (ms) by day"),
            use_container_width=True,
        )

        # Per-model analytics
        st.markdown("### Per-model analytics")
        m1, m2 = st.columns(2)
        m1.plotly_chart(
            px.bar(
                df.fillna({"model_name": "(unknown)"}).groupby("model_name").size().reset_index(name="count"),
                x="model_name",
                y="count",
                title="Interactions by model",
            ),
            use_container_width=True,
        )
        m2.plotly_chart(
            px.bar(
                df.fillna({"model_name": "(unknown)"}).groupby("model_name")["latency_ms"].mean().reset_index(),
                x="model_name",
                y="latency_ms",
                title="Avg latency (ms) by model",
            ),
            use_container_width=True,
        )

        # Satisfaction
        if "rating" in df.columns:
            st.markdown("### Satisfaction (👍=1, 👎=-1)")
            sat = df.dropna(subset=["rating"]).copy()
            if len(sat) > 0:
                sat["pos"] = (sat["rating"] > 0).astype(int)
                s1, s2 = st.columns(2)
                s1.plotly_chart(
                    px.bar(
                        sat.groupby("date")["pos"].mean().reset_index().rename(columns={"pos": "positive_rate"}),
                        x="date",
                        y="positive_rate",
                        title="Positive rate by day",
                    ),
                    use_container_width=True,
                )
                s2.plotly_chart(
                    px.bar(
                        sat.groupby("structure_kind")["pos"].mean().reset_index().rename(columns={"pos": "positive_rate"}),
                        x="structure_kind",
                        y="positive_rate",
                        title="Positive rate by structure",
                    ),
                    use_container_width=True,
                )
            else:
                st.info("No ratings yet. Use 👍/👎 after generating.")

# -----------------------------------------------------------------------
# SECTION: Export BI (hidden by CSS when PAIE_HIDE_BI=1)
# -----------------------------------------------------------------------
# (Tab intentionally omitted in this delivery; CSS above hides a 4th slot)




