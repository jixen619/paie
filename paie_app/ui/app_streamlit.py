# =======================================================================
# PAIE – Offline Personalized AI + Analytics
# Streamlit UI (single-file)
# =======================================================================

# ===== SECTION: Imports & Globals ======================================
import os, sys, json, subprocess, requests, sqlite3
from pathlib import Path
import yaml, streamlit as st, pandas as pd, plotly.express as px

# >>> MUST be the first Streamlit call (only once) <<<
st.set_page_config(page_title="PAIE", layout="wide")

# Optional libs (DOCX/PDF exports)
HAS_DOCX = True
try:
    from docx import Document
except Exception:
    HAS_DOCX = False
HAS_REPORTLAB = True
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
except Exception:
    HAS_REPORTLAB = False

# Hide Export BI tab if env set
if os.getenv("PAIE_HIDE_BI", "1") == "1":
    st.markdown(
        "<style>div[data-baseweb='tab-list'] button:nth-of-type(4){display:none!important;}</style>",
        unsafe_allow_html=True,
    )

# Ensure repo root on sys.path (this file is .../paie_app/ui/app_streamlit.py)
REPO_ROOT = Path(__file__).resolve().parents[2]
if str(REPO_ROOT) not in sys.path:
    sys.path.append(str(REPO_ROOT))

# Import the router from the package (consistent with your layout)
from paie_app.core.router_rag import Router  # noqa: E402

# Resolve config/db whether they live at root or under paie_app/config
ROOT_SETTINGS = REPO_ROOT / "settings.yaml"
PKG_SETTINGS  = REPO_ROOT / "paie_app" / "config" / "settings.yaml"
SETTINGS      = ROOT_SETTINGS if ROOT_SETTINGS.exists() else PKG_SETTINGS

ROOT_PROFILES = REPO_ROOT / "config" / "profiles"
PKG_PROFILES  = REPO_ROOT / "paie_app" / "config" / "profiles"
PROFILES_DIR  = ROOT_PROFILES if ROOT_PROFILES.exists() else PKG_PROFILES

ROOT_TMPL     = REPO_ROOT / "config" / "templates"
PKG_TMPL      = REPO_ROOT / "paie_app" / "config" / "templates"
TEMPLATES_DIR = ROOT_TMPL if ROOT_TMPL.exists() else PKG_TMPL

ROOT_DB = REPO_ROOT / "paie.db"
PKG_DB  = REPO_ROOT / "paie_app" / "paie.db"
DB_PATH = ROOT_DB if ROOT_DB.exists() else PKG_DB

USERNAME = "default"
HEADLINE = "PAIE – Personalized AI Engine"
st.title(HEADLINE)  # <-- keep, but DO NOT call set_page_config again

# ===== SECTION: Session Keys ===========================================
if "prompt_input" not in st.session_state:
    st.session_state["prompt_input"] = ""
if "__last_prompt" not in st.session_state:
    st.session_state["__last_prompt"] = ""

# ===== SECTION: Helpers ================================================
def list_ollama_models(host: str) -> list[str]:
    names = []
    try:
        r = requests.get(host.rstrip("/") + "/api/tags", timeout=3)
        r.raise_for_status()
        for m in (r.json() or {}).get("models", []):
            name = (m.get("name") or m.get("model") or "").strip()
            if name:
                names.append(name)
    except Exception:
        try:
            out = subprocess.check_output(["ollama", "list"], text=True, timeout=4)
            lines = [l.strip() for l in out.splitlines() if l.strip()]
            names = [ln.split()[0] for ln in lines if not ln.upper().startswith("NAME")]
        except Exception:
            pass
    seen, out = set(), []
    for n in names:
        if n not in seen:
            out.append(n); seen.add(n)
    return out

def load_profiles() -> dict:
    items = {}
    if PROFILES_DIR and PROFILES_DIR.exists():
        for p in PROFILES_DIR.glob("*.y*ml"):
            try:
                items[p.stem] = yaml.safe_load(p.read_text(encoding="utf-8")) or {}
            except Exception:
                pass
    if not items:
        items["default"] = {
            "name": "default",
            "tone": "neutral",
            "verbosity": "medium",
            "persona": "Helpful assistant",
            "style_guide": [],
            "domain_glossary": [],
        }
    return items

def load_templates() -> dict[str, str]:
    items = {"(none)": ""}
    if TEMPLATES_DIR and TEMPLATES_DIR.exists():
        for p in TEMPLATES_DIR.glob("*.txt"):
            try:
                items[p.stem] = p.read_text(encoding="utf-8")
            except Exception:
                pass
    return items

def ensure_columns():
    con = sqlite3.connect(str(DB_PATH))
    cur = con.cursor()
    for col, typ in [
        ("model_name", "TEXT"),
        ("profile_name", "TEXT"),
        ("template_name", "TEXT"),
        ("rating", "INTEGER"),
    ]:
        try:
            cur.execute(f"ALTER TABLE interactions ADD COLUMN {col} {typ}")
            con.commit()
        except Exception:
            pass
    con.close()

def update_last_interaction_meta(username, model, profile_name, template_name):
    con = sqlite3.connect(str(DB_PATH))
    cur = con.cursor()
    row = cur.execute(
        """
        SELECT i.id FROM interactions i
        JOIN users u ON u.id = i.user_id
        WHERE u.username=? ORDER BY i.id DESC LIMIT 1
        """,
        (username,),
    ).fetchone()
    if row:
        cur.execute(
            "UPDATE interactions SET model_name=?, profile_name=?, template_name=? WHERE id=?",
            (model, profile_name, template_name, row[0]),
        )
        con.commit()
    con.close()

def set_last_rating(username, val):
    con = sqlite3.connect(str(DB_PATH))
    cur = con.cursor()
    row = cur.execute(
        """
        SELECT i.id FROM interactions i
        JOIN users u ON u.id = i.user_id
        WHERE u.username=? ORDER BY i.id DESC LIMIT 1
        """,
        (username,),
    ).fetchone()
    if row:
        cur.execute("UPDATE interactions SET rating=? WHERE id=?", (val, row[0]))
        con.commit()
    con.close()
    return row[0] if row else None

def build_profile_text(p: dict) -> str:
    """Turn YAML profile into a compact system hint (robust to non-string list items)."""
    def _join_maybe_list(x):
        if isinstance(x, list):
            return "; ".join(str(i) for i in x)
        return str(x)

    lines = []
    if p.get("persona"):
        lines.append(f"Persona: {p['persona']}")
    if p.get("tone"):
        lines.append(f"Tone: {p['tone']}")
    if p.get("verbosity"):
        lines.append(f"Verbosity: {p['verbosity']}")
    if p.get("style_guide"):
        lines.append("Style guide: " + _join_maybe_list(p["style_guide"]))
    if p.get("domain_glossary"):
        lines.append("Domain glossary: " + _join_maybe_list(p["domain_glossary"]))
    return "[PROFILE]\n" + "\n".join(lines) + "\n[/PROFILE]\n"


    def _to_strings(obj):
        """Normalize obj -> list[str]. Handles str, dict, list/tuple of str/dict."""
        if obj is None:
            return []

        # if it's already a string, wrap it
        if isinstance(obj, str):
            return [obj]

        # dict -> ["k: v", ...]
        if isinstance(obj, dict):
            return [f"{k}: {v}" for k, v in obj.items()]

        # iterable of mixed items
        out = []
        if isinstance(obj, (list, tuple, set)):
            for item in obj:
                if isinstance(item, str):
                    out.append(item)
                elif isinstance(item, dict):
                    for k, v in item.items():
                        out.append(f"{k}: {v}")
                else:
                    out.append(str(item))
            return out

        # fallback
        return [str(obj)]

    lines = []
    if p.get("persona"):
        lines.append(f"Persona: {p['persona']}")
    if p.get("tone"):
        lines.append(f"Tone: {p['tone']}")
    if p.get("verbosity"):
        lines.append(f"Verbosity: {p['verbosity']}")

    style_items = _to_strings(p.get("style_guide"))
    if style_items:
        lines.append("Style guide: " + "; ".join(style_items))

    glossary_items = _to_strings(p.get("domain_glossary"))
    if glossary_items:
        lines.append("Domain glossary: " + "; ".join(glossary_items))

    return "[PROFILE]\n" + "\n".join(lines) + "\n[/PROFILE]\n"


def make_docx(prompt, response, model, structure):
    if not HAS_DOCX:
        return None
    from io import BytesIO
    doc = Document()
    doc.add_heading("PAIE Output", level=1)
    doc.add_paragraph(f"Model: {model}")
    doc.add_paragraph(f"Structure: {structure or '(none)'}")
    doc.add_heading("Prompt", level=2);   doc.add_paragraph(prompt)
    doc.add_heading("Response", level=2); doc.add_paragraph(response)
    buf = BytesIO(); doc.save(buf); buf.seek(0); return buf.getvalue()

def make_pdf(text):
    if not HAS_REPORTLAB:
        return None
    from io import BytesIO
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4
    x, y = 40, height - 40
    for line in text.splitlines():
        if y < 40:
            c.showPage(); y = height - 40
        c.drawString(x, y, line[:120]); y -= 14
    c.save(); buf.seek(0); return buf.getvalue()


def apply_custom_style(
    base_size_px: int,
    output_size_px: int,
    code_size_px: int,
    family_css: str,
    line_height: float,
    accent_hex: str,
):
    """
    Injects CSS that tweaks global fonts, text areas, code blocks, and a
    .paie-output wrapper used for the model response.
    """
    css = f"""
    <style>
      /* Global font & base size */
      html, body, [data-testid="stAppViewContainer"] * {{
        font-family: {family_css}, -apple-system, "Segoe UI", Roboto, "Helvetica Neue",
                      Arial, "Noto Sans", "Liberation Sans", sans-serif;
        font-size: {base_size_px}px;
        line-height: {line_height};
      }}

      /* Prompt textarea */
      .stTextArea textarea {{
        font-size: {base_size_px}px !important;
        line-height: {line_height} !important;
      }}

      /* Model output wrapper */
      .paie-output {{
        font-size: {output_size_px}px;
        line-height: {line_height};
      }}
      .paie-output h1 {{ font-size: calc({output_size_px}px * 1.6); margin-top: .8em; }}
      .paie-output h2 {{ font-size: calc({output_size_px}px * 1.35); margin-top: .8em; }}
      .paie-output h3 {{ font-size: calc({output_size_px}px * 1.2);  margin-top: .8em; }}

      /* Code blocks */
      pre, code, .stCode, .stMarkdown pre code {{
        font-size: {code_size_px}px !important;
      }}

      /* Buttons accent color */
      :root {{ --paie-accent: {accent_hex}; }}
      .stButton>button, .stDownloadButton>button {{
        background: var(--paie-accent);
        border: 1px solid var(--paie-accent);
        border-radius: 12px;
      }}
    </style>
    """
    st.markdown(css, unsafe_allow_html=True)

# ===== SECTION: Sidebar – Session ======================================
st.sidebar.header("Session")
with open(SETTINGS, "r", encoding="utf-8") as f:
    default_settings = yaml.safe_load(f) or {}

available_models = list_ollama_models((default_settings.get("model") or {}).get("host", "http://localhost:11434"))
default_model    = (default_settings.get("model") or {}).get("name", "llama3.2:latest")
if not available_models:
    available_models = [default_model]
if default_model in available_models:
    available_models = [default_model] + [m for m in available_models if m != default_model]

model = st.sidebar.selectbox("Model", options=available_models, index=0)
kind  = st.sidebar.selectbox("Structure", ["(none)", "user_story", "use_case", "test_case", "summary"], index=0)

profiles      = load_profiles()
profile_names = list(profiles.keys())
profile_name  = st.sidebar.selectbox(
    "Profile", options=profile_names,
    index=profile_names.index("default") if "default" in profile_names else 0
)

templates     = load_templates()
template_name = st.sidebar.selectbox("Template", options=list(templates.keys()), index=0)

# RAG & sampling controls
use_rag     = st.sidebar.checkbox("Use RAG", value=True)
rag_top_k   = st.sidebar.slider("RAG Top-K", 1, 8, value=int((default_settings.get("rag") or {}).get("top_k", 4)))
temperature = st.sidebar.slider("Temperature", 0.0, 1.5, value=float((default_settings.get("model") or {}).get("temperature", 0.3)), step=0.05)
top_p       = st.sidebar.slider("Top-p",       0.0, 1.0, value=float((default_settings.get("model") or {}).get("top_p",       0.95)), step=0.01)

# Init DB + router
ensure_columns()
r = Router(str(SETTINGS))
r.settings["model"]["name"] = model  # propagate chosen model

# ===== SECTION: Tabs ====================================================
tabs = st.tabs(["Generate", "History", "Analytics"])

# --------------------------- SECTION: Generate --------------------------
with tabs[0]:
    prompt = st.text_area("Prompt", key="prompt_input", height=160,
                          placeholder="Ask for a User Story / Use Case / Test Case ...")

    left, spacer, right = st.columns([1, 6, 1])

    def _clear_prompt():
        st.session_state["prompt_input"] = ""
        st.session_state["__last_prompt"] = ""

    with left:
        run_clicked = st.button("Generate", type="primary", key="btn_generate")
    with right:
        st.button("Clear input", key="btn_clear", on_click=_clear_prompt)

    # RAG sources (preview)
    with st.expander("RAG sources (preview)"):
        try:
            from rag import retrieve
            q = (st.session_state.get("prompt_input") or "").strip()
            if q:
                hits = retrieve(q, k=rag_top_k)
                if not hits:
                    st.caption("No hits (empty index or non-text docs).")
                for i, (text, meta) in enumerate(hits, 1):
                    st.markdown(f"**{i}. {(meta or {}).get('source','')}**")
                    st.code((text or "")[:600])
            else:
                st.caption("Type a prompt to preview matches.")
        except Exception as e:
            st.caption(f"RAG preview unavailable: {e}")

    if run_clicked and st.session_state.get("prompt_input", "").strip():
        prompt = st.session_state["prompt_input"]
        prof_txt = build_profile_text(profiles[profile_name])
        tpl_txt  = templates.get(template_name, "")
        final_prompt = f"{prof_txt}{('[TEMPLATE]\\n'+tpl_txt+'\\n[/TEMPLATE]\\n') if tpl_txt else ''}{prompt}"

        with st.spinner("Thinking locally..."):
            k_tag = None if kind == "(none)" else kind
            out, ms = r.run(
                username=USERNAME,
                prompt=final_prompt,
                kind=k_tag,
                temperature=temperature,
                top_p=top_p,
                use_rag=use_rag,
                rag_top_k=rag_top_k,
            )

        update_last_interaction_meta(USERNAME, model, profile_name, template_name)
        st.success(f"Done in {ms} ms")
        st.markdown(out)

        # Downloads
        md = f"# Model: {model}\n# Structure: {k_tag or '(none)'}\n\n# Prompt\n\n{prompt}\n\n# Response\n\n{out}"
        st.download_button("Download Markdown", data=md, file_name="output.md", mime="text/markdown")

        if HAS_DOCX:
            docx_bytes = make_docx(prompt, out, model, k_tag)
            if docx_bytes:
                st.download_button(
                    "Download DOCX",
                    data=docx_bytes,
                    file_name="output.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
        else:
            st.info("DOCX export requires python-docx (optional).")

        if HAS_REPORTLAB:
            pdf_bytes = make_pdf(f"Model: {model}\nStructure: {k_tag or '(none)'}\n\nPrompt:\n{prompt}\n\nResponse:\n{out}")
            if pdf_bytes:
                st.download_button("Download PDF", data=pdf_bytes, file_name="output.pdf", mime="application/pdf")
        else:
            st.info("PDF export requires reportlab (optional).")

        st.write("Feedback:")
        b1, b2 = st.columns(2)
        if b1.button("👍 Good"):
            st.toast(f"Thanks! Saved 👍 for interaction #{set_last_rating(USERNAME, 1)}")
        if b2.button("👎 Needs work"):
            st.toast(f"Saved 👎 for interaction #{set_last_rating(USERNAME, -1)}")

# ---------------------------- SECTION: History --------------------------
with tabs[1]:
    st.subheader("History")
    con = sqlite3.connect(str(DB_PATH))
    q = st.text_input("Search", placeholder="filter by keyword (optional)")

    sql = """
      SELECT i.id, i.created_at, i.structure_kind, i.prompt, i.response,
             i.model_name, i.profile_name, i.template_name, i.rating
      FROM interactions i JOIN users u ON u.id=i.user_id
      WHERE u.username=?
    """
    params = [USERNAME]
    if q:
        sql += " AND (i.prompt LIKE ? OR i.response LIKE ?)"
        params += [f"%{q}%", f"%{q}%"]
    sql += " ORDER BY i.id DESC LIMIT 200"
    df = pd.read_sql_query(sql, con, params=params)

    with st.expander("🧹 Clear history"):
        st.caption("Permanently deletes saved prompts/responses.")
        if st.button("Delete all history"):
            uid = con.execute("SELECT id FROM users WHERE username=?", (USERNAME,)).fetchone()
            if uid:
                con.execute("DELETE FROM interactions WHERE user_id=?", (uid[0],))
                con.commit()
                st.success("History cleared. Click ↻ Rerun.")
            else:
                st.info("No records found.")
    con.close()

    if df.empty:
        st.info("No history yet.")
    else:
        df["date"] = pd.to_datetime(df["created_at"]).dt.date
        for date, grp in df.groupby("date", sort=False):
            st.markdown(f"### {date}")
            for _, row in grp.iterrows():
                header = f"#{row['id']}  •  {row['created_at']}  •  {row['structure_kind'] or '(none)'}  •  {row['model_name'] or '(model)'}"
                with st.expander(header):
                    st.markdown(
                        f"**Profile**: {row['profile_name'] or profile_name}  |  "
                        f"**Template**: {row['template_name'] or '(none)'}  |  "
                        f"**Rating**: {row['rating'] if pd.notna(row['rating']) else '—'}"
                    )
                    st.markdown("**Prompt**");   st.code(row["prompt"])
                    st.markdown("**Response**"); st.markdown(row["response"])

                    md_hist = f"# Prompt\n\n{row['prompt']}\n\n# Response\n\n{row['response']}"
                    st.download_button(
                        "Download as Markdown",
                        data=md_hist,
                        file_name=f"paie_{row['id']}.md",
                        key=f"md-{row['id']}",
                    )
                    if HAS_DOCX:
                        docx_bytes = make_docx(
                            row["prompt"], row["response"], row["model_name"] or "", row["structure_kind"]
                        )
                        if docx_bytes:
                            st.download_button(
                                "Download as DOCX",
                                data=docx_bytes,
                                file_name=f"paie_{row['id']}.docx",
                                key=f"docx-{row['id']}",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            )

# --------------------------- SECTION: Analytics -------------------------
with tabs[2]:
    st.subheader("Usage & Performance (from SQLite)")
    con = sqlite3.connect(str(DB_PATH))
    df = pd.read_sql_query(
        "SELECT created_at, structure_kind, latency_ms, model_name, rating FROM interactions ORDER BY id DESC LIMIT 2000",
        con,
    )
    con.close()

    if len(df) == 0:
        st.info("No data yet. Generate something first.")
    else:
        df["date"] = pd.to_datetime(df["created_at"]).dt.date
        c1, c2 = st.columns(2)
        c1.plotly_chart(px.bar(df, x="structure_kind", title="Structure usage count"), use_container_width=True)
        c2.plotly_chart(
            px.line(df.groupby("date")["latency_ms"].mean().reset_index(), x="date", y="latency_ms", title="Avg latency (ms) by day"),
            use_container_width=True,
        )

        st.markdown("### Per-model analytics")
        m1, m2 = st.columns(2)
        m1.plotly_chart(
            px.bar(
                df.fillna({"model_name": "(unknown)"}).groupby("model_name").size().reset_index(name="count"),
                x="model_name",
                y="count",
                title="Interactions by model",
            ),
            use_container_width=True,
        )
        m2.plotly_chart(
            px.bar(
                df.fillna({"model_name": "(unknown)"}).groupby("model_name")["latency_ms"].mean().reset_index(),
                x="model_name",
                y="latency_ms",
                title="Avg latency (ms) by model",
            ),
            use_container_width=True,
        )

        if "rating" in df.columns:
            st.markdown("### Satisfaction (👍=1, 👎=-1)")
            sat = df.dropna(subset=["rating"]).copy()
            if len(sat) > 0:
                sat["pos"] = (sat["rating"] > 0).astype(int)
                s1, s2 = st.columns(2)
                s1.plotly_chart(
                    px.bar(
                        sat.groupby("date")["pos"].mean().reset_index().rename(columns={"pos": "positive_rate"}),
                        x="date",
                        y="positive_rate",
                        title="Positive rate by day",
                    ),
                    use_container_width=True,
                )
                s2.plotly_chart(
                    px.bar(
                        sat.groupby("structure_kind")["pos"].mean().reset_index().rename(columns={"pos": "positive_rate"}),
                        x="structure_kind",
                        y="positive_rate",
                        title="Positive rate by structure",
                    ),
                    use_container_width=True,
                )
            else:
                st.info("No ratings yet. Use 👍/👎 after generating.")
